from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID, uuid4

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

import app.report_export as report_export
from app.main import app
from app.report_export import ReportExportError, export_audit_report
from app.report_export_models import (
    AuditReportDocument,
    ReportContractEvidence,
    ReportExportFormat,
    ReportHumanDecision,
    ReportIssue,
    ReportLegalEvidence,
)
from app.storage_management import delete_job_storage


def _sample_report(job_id: UUID) -> AuditReportDocument:
    return AuditReportDocument(
        job_id=job_id,
        filename="测试劳动合同.docx",
        document_kind="docx",
        as_of="2026-08-21",
        overall_state="HUMAN_REVIEW_REQUIRED",
        contract_type="EMPLOYMENT",
        planning_mode="DIRECT",
        planning_coverage_complete=True,
        canonical_object_count=8,
        reviewed_with_issue_count=1,
        reviewed_no_specific_issue_count=7,
        primary_provider="deepseek",
        primary_model="deepseek-v4-pro",
        secondary_provider="kimi",
        secondary_model="kimi-k3",
        final_review_state="HUMAN_REVIEW_REQUIRED",
        human_review_required_count=1,
        human_review_resolved_required_count=1,
        human_review_outstanding_required_count=0,
        issues=[
            ReportIssue(
                issue_id="ISSUE-001",
                topic="解除劳动合同经济补偿",
                priority="HIGH_ATTENTION",
                why_review=["解除条款可能影响劳动者法定补偿权利。"],
                questions=["补偿标准是否低于法定标准？"],
                primary_state="SUPPORTED_FINDING",
                primary_severity="HIGH",
                primary_title="经济补偿约定可能低于法定标准",
                primary_reasoning="合同约定与法定经济补偿计算规则存在冲突风险。",
                primary_suggestion="按现行劳动合同法重新核对并修改补偿条款。",
                primary_evidence_sufficiency="SUFFICIENT",
                secondary_assessment="SUPPORTED",
                secondary_coverage_assessment="COVERED",
                secondary_severity="HIGH",
                secondary_reasoning="复核认为主要风险判断有合同证据和法律依据支持。",
                secondary_suggestion="保留人工复核并修订相关条款。",
                comparison_state="CONSISTENT_WITH_REVIEW",
                requires_human_review=True,
                comparison_reasons=["双模型结论一致，但高风险事项仍要求人工确认。"],
                contract_evidence=[
                    ReportContractEvidence(
                        evidence_id="EVID-001",
                        quote="甲方解除劳动合同时仅支付固定补偿1000元。",
                        source_method="DOCX_NATIVE_TEXT",
                    )
                ],
                legal_evidence=[
                    ReportLegalEvidence(
                        legal_evidence_id="LAW-001",
                        authority_id="prc-labor-contract-law",
                        authority_title="中华人民共和国劳动合同法",
                        version_id="effective-2012-12-28",
                        article_token="第四十七条",
                        article_text="经济补偿按劳动者在本单位工作的年限计算。",
                        effective_date="2013-07-01",
                        coverage_type="FULL_TEXT",
                    )
                ],
                human_decision=ReportHumanDecision(
                    state="CONFIRMED",
                    revision=1,
                    decided_at=datetime(2026, 8, 21, 9, 0, tzinfo=timezone.utc),
                    reviewer_note="确认需要修改。",
                    is_stale=False,
                ),
            )
        ],
        source_uncertainty=[],
        warnings=["报告只反映本地语料和已完成审计链。"],
        source_fingerprints={
            "contract_source": "a" * 64,
            "issue_review_report": "b" * 64,
        },
        report_content_fingerprint="c" * 64,
    )


def test_docx_and_pdf_render_to_openable_files(tmp_path: Path) -> None:
    report = _sample_report(uuid4())
    docx_path = tmp_path / "report.docx"
    pdf_path = tmp_path / "report.pdf"

    report_export._render_docx(report, docx_path)
    report_export._render_pdf(report, pdf_path)

    opened = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in opened.paragraphs)
    assert "Law-Rag 合同审计报告" in text
    assert "解除劳动合同经济补偿" in text
    assert "中华人民共和国劳动合同法" in text
    assert "确认需要修改" in text

    reader = PdfReader(str(pdf_path))
    assert len(reader.pages) >= 1
    assert pdf_path.read_bytes().startswith(b"%PDF")


def test_export_writes_only_job_owned_export_and_hash_manifest(tmp_path: Path, monkeypatch) -> None:
    job_id = uuid4()
    report = _sample_report(job_id)
    monkeypatch.setenv("LAW_RAG_RUNTIME_DIR", str(tmp_path))
    monkeypatch.setattr(report_export, "build_audit_report", lambda value: report if value == job_id else None)

    path, result = export_audit_report(job_id, ReportExportFormat.DOCX)

    assert path.parent == tmp_path / "exports" / str(job_id)
    assert path.is_file()
    assert result.size_bytes == path.stat().st_size
    assert len(result.sha256) == 64
    manifest = path.with_name(path.name + ".manifest.json")
    payload = json.loads(manifest.read_text(encoding="utf-8"))
    assert payload["sha256"] == result.sha256
    assert "测试劳动合同" not in manifest.read_text(encoding="utf-8")
    assert not (tmp_path / "legal").exists()


def test_export_root_symlink_fails_closed(tmp_path: Path, monkeypatch) -> None:
    job_id = uuid4()
    outside = tmp_path / "outside"
    outside.mkdir()
    (tmp_path / "exports").symlink_to(outside, target_is_directory=True)
    monkeypatch.setenv("LAW_RAG_RUNTIME_DIR", str(tmp_path))
    monkeypatch.setattr(report_export, "build_audit_report", lambda _: _sample_report(job_id))

    with pytest.raises(ReportExportError, match="symlink"):
        export_audit_report(job_id, ReportExportFormat.PDF)

    assert not any(outside.iterdir())


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _terminal_job(root: Path, job_id: UUID) -> None:
    updated = "2026-08-21T09:00:00+00:00"
    job_dir = root / "jobs" / str(job_id)
    _write_json(job_dir / "document.json", {"job_id": str(job_id), "filename": "contract.pdf", "document_kind": "pdf"})
    _write_json(
        job_dir / "pipeline.json",
        {
            "schema_version": "1.3.0",
            "engine_version": "stage13g-4-1.0.0",
            "job_id": str(job_id),
            "status": "COMPLETE",
            "current_stage": "COMPLETE",
            "progress_percent": 100,
            "as_of": "2026-08-21",
            "use_semantic": False,
            "started_at": "2026-08-21T08:00:00+00:00",
            "updated_at": updated,
            "completed_at": updated,
            "failure_code": None,
            "failure_detail": None,
            "stages": [],
        },
    )
    upload = root / "uploads" / str(job_id) / "source.pdf"
    upload.parent.mkdir(parents=True, exist_ok=True)
    upload.write_bytes(b"PDF")


def test_job_cleanup_reclaims_exports_and_never_touches_shared_legal(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("LAW_RAG_RUNTIME_DIR", str(tmp_path))
    job_id = uuid4()
    _terminal_job(tmp_path, job_id)
    export = tmp_path / "exports" / str(job_id) / "audit.pdf"
    export.parent.mkdir(parents=True, exist_ok=True)
    export.write_bytes(b"PRIVATE-REPORT")
    legal = tmp_path / "legal" / "legal.db"
    legal.parent.mkdir(parents=True, exist_ok=True)
    legal.write_bytes(b"SHARED-LEGAL")

    result = delete_job_storage(job_id, confirm_job_id=job_id)

    assert result.deleted is True
    assert not export.parent.exists()
    assert legal.read_bytes() == b"SHARED-LEGAL"
    assert result.reclaimed_bytes >= len(b"PRIVATE-REPORT")


def test_report_export_api_returns_attachment_headers(tmp_path: Path, monkeypatch) -> None:
    job_id = uuid4()
    output = tmp_path / "report.pdf"
    output.write_bytes(b"%PDF-test")

    from app import report_export_api
    from app.report_export_models import ReportExportResult

    result = ReportExportResult(
        job_id=job_id,
        format=ReportExportFormat.PDF,
        filename="audit.pdf",
        size_bytes=output.stat().st_size,
        sha256="d" * 64,
        report_content_fingerprint="e" * 64,
    )
    monkeypatch.setattr(report_export_api, "export_audit_report", lambda *_: (output, result))
    response = TestClient(app).post(f"/api/documents/{job_id}/report-export/pdf")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.headers["x-law-rag-report-sha256"] == "d" * 64
    assert response.headers["x-law-rag-report-content-fingerprint"] == "e" * 64
