from __future__ import annotations

import hashlib
from pathlib import Path
from uuid import uuid4

from docx import Document
from pypdf import PdfReader
from reportlab.pdfgen import canvas

from app import tester_report_watermark
from app.report_export_models import ReportExportFormat, ReportExportResult
from app.tester_report_watermark import apply_tester_report_watermark


WATERMARK = "Law-Rag 0.8.0-rc3-tester1 · Tester T001 · Limited Test Build"


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _result(path: Path, format: ReportExportFormat) -> ReportExportResult:
    return ReportExportResult(
        job_id=uuid4(),
        format=format,
        filename=path.name,
        size_bytes=path.stat().st_size,
        sha256=_sha256(path),
        report_content_fingerprint="a" * 64,
    )


def test_docx_footer_gets_tester_id_and_hash_is_rebound(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Audit body")
    document.save(path)
    original = _result(path, ReportExportFormat.DOCX)
    monkeypatch.setattr(tester_report_watermark, "active_tester_watermark", lambda: WATERMARK)

    updated = apply_tester_report_watermark(path, original)

    reopened = Document(path)
    footer_text = "\n".join(paragraph.text for section in reopened.sections for paragraph in section.footer.paragraphs)
    assert WATERMARK in footer_text
    assert updated.sha256 == _sha256(path)
    assert updated.sha256 != original.sha256
    assert updated.report_content_fingerprint == original.report_content_fingerprint
    assert (tmp_path / "report.docx.manifest.json").is_file()


def test_pdf_every_page_gets_tester_id_and_hash_is_rebound(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "report.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 720, "Page one")
    pdf.showPage()
    pdf.drawString(72, 720, "Page two")
    pdf.save()
    original = _result(path, ReportExportFormat.PDF)
    monkeypatch.setattr(tester_report_watermark, "active_tester_watermark", lambda: WATERMARK)

    updated = apply_tester_report_watermark(path, original)

    reader = PdfReader(str(path))
    assert len(reader.pages) == 2
    for page in reader.pages:
        extracted = page.extract_text() or ""
        assert "Tester T001" in extracted
        assert "Limited Test Build" in extracted
    assert updated.sha256 == _sha256(path)
    assert updated.sha256 != original.sha256
    assert updated.report_content_fingerprint == original.report_content_fingerprint
    assert (tmp_path / "report.pdf.manifest.json").is_file()
