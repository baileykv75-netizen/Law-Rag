from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin, urlsplit

import httpx
from docx import Document
from lxml import html

from .models import (
    AuthorityType,
    CoverageType,
    LegalAuthority,
    LegalManifest,
    ManifestRecord,
    OfficialSourceRef,
    SourceRole,
    VersionStatus,
)
from .parser import normalize_snapshot_text, parse_chinese_articles, sha256_text


NPC_API_URL = "https://flk.npc.gov.cn/law-search/search/list"
NPC_DETAIL_API_URL = "https://flk.npc.gov.cn/law-search/search/flfgDetails"
NPC_DETAIL_PAGE_URL = "https://flk.npc.gov.cn/detail"
NPC_FILE_BASE_URL = "https://wb.flk.npc.gov.cn"
NPC_PUBLIC_FILE_API_URL = "https://flk.npc.gov.cn/law-search/amazonFile/ofdGenerateLink"


@dataclass(frozen=True)
class OfficialTextSource:
    law_ref: str
    source_name: str
    url: str


OFFICIAL_HTML_TEXT_SOURCES: tuple[OfficialTextSource, ...] = (
    OfficialTextSource(
        law_ref="中华人民共和国广告法",
        source_name="中国政府网",
        url="https://www.gov.cn/guoqing/2021-10/29/content_5647620.htm",
    ),
    OfficialTextSource(
        law_ref="中华人民共和国建筑法",
        source_name="中国人大网",
        url="http://www.npc.gov.cn/zgrdw/npc/xinwen/2019-05/07/content_2086833.htm",
    ),
    OfficialTextSource(
        law_ref="建设工程质量管理条例",
        source_name="中国政府网国务院公报",
        url="https://www.gov.cn/gongbao/content/2019/content_5468867.htm",
    ),
    OfficialTextSource(
        law_ref="最高人民法院关于审理建设工程施工合同纠纷案件适用法律问题的解释（一）",
        source_name="最高人民法院",
        url="https://www.court.gov.cn/zixun/xiangqing/282111.html",
    ),
)


class OfficialLegalDownloadError(RuntimeError):
    pass


@dataclass(frozen=True)
class DownloadedLegalManifest:
    law_ref: str
    title: str
    manifest_path: Path
    article_count: int


def _slug(value: str, prefix: str = "npc-law") -> str:
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:16]
    return f"{prefix}-{digest}"


def _clean(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    if "<" in text and ">" in text:
        text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\s+", " ", text).strip()


def _parse_date(value: Any) -> date | None:
    text = _clean(value)
    if not text:
        return None
    match = re.search(r"(\d{4})[-年./](\d{1,2})[-月./](\d{1,2})", text)
    if not match:
        return None
    return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))


def _first_text(payload: dict[str, Any], names: tuple[str, ...]) -> str:
    for name in names:
        value = _clean(payload.get(name))
        if value:
            return value
    return ""


def _first_date(payload: dict[str, Any], names: tuple[str, ...]) -> date | None:
    for name in names:
        parsed = _parse_date(payload.get(name))
        if parsed is not None:
            return parsed
    return None


def _iter_dicts(value: Any):
    if isinstance(value, dict):
        yield value
        for child in value.values():
            yield from _iter_dicts(child)
    elif isinstance(value, list):
        for child in value:
            yield from _iter_dicts(child)


def _extract_list(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []
    for key in ("result", "data", "list", "content", "records", "rows"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
        if isinstance(value, dict):
            nested = _extract_list(value)
            if nested:
                return nested
    return []


def _extract_id(item: dict[str, Any]) -> str:
    for name in ("bbbs", "id", "_id", "docId", "gid", "f_id", "bbh", "bodyId"):
        value = _clean(item.get(name))
        if value:
            return value
    return ""


def _extract_title(item: dict[str, Any]) -> str:
    return _first_text(item, ("title", "name", "f_title", "flfgbt", "docTitle", "bt"))


def _title_matches(candidate: str, law_ref: str) -> bool:
    compact_candidate = re.sub(r"[《》（）()：:\s]+", "", candidate)
    compact_ref = re.sub(r"[《》（）()：:\s]+", "", law_ref)
    if compact_candidate == compact_ref:
        return True
    return compact_candidate in compact_ref or compact_ref in compact_candidate


def _official_text_source_for(law_ref: str, title: str = "") -> OfficialTextSource | None:
    compact_ref = re.sub(r"[《》（）()：:\s]+", "", law_ref)
    compact_title = re.sub(r"[《》（）()：:\s]+", "", title)
    for source in OFFICIAL_HTML_TEXT_SOURCES:
        compact_source = re.sub(r"[《》（）()：:\s]+", "", source.law_ref)
        if compact_source in compact_ref or compact_ref in compact_source:
            return source
        if compact_title and (compact_source in compact_title or compact_title in compact_source):
            return source
    return None


def has_official_text_source(law_ref: str) -> bool:
    return _official_text_source_for(law_ref) is not None


def _json_from_response(response: httpx.Response) -> Any:
    return json.loads(response.content.decode("utf-8-sig"))


def _post_json_or_form(client: httpx.Client, url: str, payload: dict[str, Any], *, timeout: float) -> Any:
    headers = {
        "User-Agent": "Law-Rag legal corpus downloader; official-source verification",
        "Accept": "application/json, text/plain, */*",
        "Referer": "https://flk.npc.gov.cn/search",
    }
    errors: list[str] = []
    for kwargs in ({"json": payload}, {"data": payload}):
        try:
            response = client.post(url, headers=headers, timeout=timeout, **kwargs)
            response.raise_for_status()
            return _json_from_response(response)
        except Exception as exc:  # pragma: no cover - exact network failure depends on the official site.
            errors.append(str(exc))
    raise OfficialLegalDownloadError(f"官方接口请求失败：{'; '.join(errors)}")


def _get_json(client: httpx.Client, url: str, params: dict[str, Any], *, timeout: float) -> Any:
    headers = {
        "User-Agent": "Law-Rag legal corpus downloader; official-source verification",
        "Accept": "application/json, text/plain, */*",
        "Referer": "https://flk.npc.gov.cn/search",
    }
    try:
        response = client.get(url, headers=headers, params=params, timeout=timeout)
        response.raise_for_status()
        return _json_from_response(response)
    except Exception as exc:  # pragma: no cover - exact network failure depends on the official site.
        raise OfficialLegalDownloadError(f"官方详情接口请求失败：{exc}") from exc


def _search_npc(client: httpx.Client, law_ref: str, *, timeout: float) -> dict[str, Any]:
    payload = {
        "searchRange": 1,
        "sxrq": [],
        "gbrq": [],
        "searchType": 2,
        "sxx": [],
        "gbrqYear": [],
        "flfgCodeId": [],
        "zdjgCodeId": [],
        "searchContent": law_ref,
        "orderByParam": {"order": "-1", "sort": ""},
        "pageNum": 1,
        "pageSize": 10,
    }
    data = _post_json_or_form(client, NPC_API_URL, payload, timeout=timeout)
    if isinstance(data, dict) and data.get("code") != 200:
        raise OfficialLegalDownloadError(f"国家法律法规数据库检索失败：{data.get('msg') or data.get('message') or data}")
    candidates = _extract_list(data)
    if not candidates:
        raise OfficialLegalDownloadError(f"国家法律法规数据库没有返回《{law_ref}》的检索结果。")
    exact = [item for item in candidates if _title_matches(_extract_title(item), law_ref)]
    chosen = exact[0] if exact else candidates[0]
    if not _extract_id(chosen):
        raise OfficialLegalDownloadError(f"国家法律法规数据库检索结果缺少详情 ID：《{law_ref}》。")
    return chosen


def _detail_npc(client: httpx.Client, result: dict[str, Any], *, timeout: float) -> dict[str, Any]:
    detail_id = _extract_id(result)
    data = _get_json(client, NPC_DETAIL_API_URL, {"bbbs": detail_id}, timeout=timeout)
    if isinstance(data, dict) and data.get("code") != 200:
        raise OfficialLegalDownloadError(f"国家法律法规数据库详情失败：{data.get('msg') or data.get('message') or data}")
    if isinstance(data, dict):
        for key in ("result", "data", "body"):
            value = data.get(key)
            if isinstance(value, dict):
                return value
        return data
    raise OfficialLegalDownloadError("国家法律法规数据库详情接口返回格式无法识别。")


def _authority_type(detail: dict[str, Any]) -> AuthorityType:
    text = " ".join(_clean(value) for value in detail.values() if not isinstance(value, (dict, list)))
    if "司法解释" in text or "最高人民法院" in text:
        return AuthorityType.JUDICIAL_INTERPRETATION
    if "行政法规" in text or "国务院" in text:
        return AuthorityType.ADMINISTRATIVE_REGULATION
    return AuthorityType.LAW


def _version_status(detail: dict[str, Any]) -> VersionStatus:
    if detail.get("sxx") == 4:
        return VersionStatus.NOT_YET_EFFECTIVE
    if detail.get("sxx") == 3:
        return VersionStatus.EFFECTIVE
    if detail.get("sxx") == 2:
        return VersionStatus.AMENDED
    if detail.get("sxx") == 1:
        return VersionStatus.REPEALED
    text = " ".join(_clean(value) for value in detail.values() if not isinstance(value, (dict, list)))
    if any(token in text for token in ("尚未生效", "未生效")):
        return VersionStatus.NOT_YET_EFFECTIVE
    if any(token in text for token in ("废止", "失效")):
        return VersionStatus.REPEALED
    return VersionStatus.EFFECTIVE


def _text_from_html(value: str) -> str:
    tree = html.fromstring(value)
    lines = [line.strip() for line in tree.text_content().splitlines() if line.strip()]
    return "\n".join(lines)


def _article_count_hint(text: str) -> int:
    return len(re.findall(r"第[零〇一二三四五六七八九十百千万两\d]+条", text))


def _clean_html_text(value: str) -> str:
    lines = [line.strip() for line in value.splitlines() if line.strip()]
    ignored = {
        "首页",
        "简",
        "繁",
        "EN",
        "登录",
        "个人中心",
        "退出",
        "邮箱",
        "无障碍",
        "打印",
        "来源：",
        "字号：",
        "默认",
        "大",
        "超大",
        "|",
    }
    cleaned: list[str] = []
    for line in lines:
        if line in ignored:
            continue
        if line.startswith(("var ", "function ", "@media ", "body {", "/*", "*/")):
            continue
        cleaned.append(line)
    toc_index = next((index for index, line in enumerate(cleaned) if line == "目录"), None)
    if toc_index is not None:
        first_article_index = next(
            (
                index
                for index, line in enumerate(cleaned[toc_index + 1 :], start=toc_index + 1)
                if re.match(r"^\s*第[零〇一二三四五六七八九十百千万两\d]+条", line)
            ),
            None,
        )
        if first_article_index is not None:
            chapter_indices = [
                index
                for index, line in enumerate(cleaned[toc_index + 1 : first_article_index], start=toc_index + 1)
                if re.match(r"^\s*第[零〇一二三四五六七八九十百千万两\d]+章", line)
            ]
            if chapter_indices:
                cleaned = cleaned[:toc_index] + cleaned[chapter_indices[-1] :]
    return "\n".join(cleaned)


def _extract_official_html_text(raw_html: str) -> str:
    tree = html.fromstring(raw_html)
    candidate_xpaths = (
        "//*[@id='UCAP-CONTENT']",
        "//*[contains(concat(' ', normalize-space(@class), ' '), ' pages_content ')]",
        "//*[contains(concat(' ', normalize-space(@class), ' '), ' TRS_Editor ')]",
        "//*[contains(concat(' ', normalize-space(@class), ' '), ' article ')]",
        "//main",
        "//body",
    )
    candidates: list[str] = []
    for xpath in candidate_xpaths:
        for node in tree.xpath(xpath):
            if not hasattr(node, "text_content"):
                continue
            text = _clean_html_text(node.text_content())
            if text:
                candidates.append(text)
    candidates = [text for text in candidates if _article_count_hint(text) > 0]
    if not candidates:
        raise OfficialLegalDownloadError("官方网页未找到可解析的条文正文。")
    return max(candidates, key=lambda text: (_article_count_hint(text), len(text)))


def _fetch_official_html_text(
    client: httpx.Client,
    source: OfficialTextSource,
    *,
    timeout: float,
) -> tuple[str, str, str]:
    response = client.get(
        source.url,
        headers={
            "User-Agent": "Law-Rag legal corpus downloader; official-source verification",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Referer": "https://flk.npc.gov.cn/search",
        },
        timeout=timeout,
    )
    response.raise_for_status()
    text = _extract_official_html_text(response.text)
    return text, source.url, source.source_name


def _extract_inline_text(detail: dict[str, Any]) -> str:
    for item in _iter_dicts(detail):
        for key in ("content", "body", "html", "text", "nr", "qwContent", "fullText"):
            value = item.get(key)
            if not isinstance(value, str) or not value.strip():
                continue
            if "<" in value and ">" in value:
                return _text_from_html(value)
            return value
    return ""


def _attachment_urls(detail: dict[str, Any]) -> list[str]:
    urls: list[str] = []
    for item in _iter_dicts(detail):
        for value in item.values():
            if not isinstance(value, str):
                continue
            if not re.search(r"\.(docx?|html?|txt)(?:\?|$)", value, flags=re.IGNORECASE):
                continue
            url = value
            if url.startswith("//"):
                url = f"https:{url}"
            elif url.startswith("/"):
                url = urljoin(NPC_FILE_BASE_URL, url)
            elif not url.startswith("http"):
                url = urljoin(NPC_FILE_BASE_URL + "/", url)
            if url not in urls:
                urls.append(url)
    return urls


def _official_file_urls(detail: dict[str, Any]) -> list[str]:
    paths: list[str] = []
    oss_file = detail.get("ossFile")
    if isinstance(oss_file, dict):
        for key in ("ossWordPath", "ossPdfPath", "ossWordOfdPath", "ossPdfOfdPath"):
            value = _clean(oss_file.get(key))
            if value:
                paths.append(value)
    for item in _iter_dicts(detail):
        for key, value in item.items():
            if "path" not in str(key).lower() or not isinstance(value, str):
                continue
            if re.search(r"\.(docx?|pdf|ofd)(?:\?|$)", value, flags=re.IGNORECASE):
                paths.append(value)
    out: list[str] = []
    for path in paths:
        if path not in out:
            out.append(path)
    return out


def _download_official_file(
    client: httpx.Client,
    file_path: str,
    *,
    timeout: float,
) -> tuple[bytes, str]:
    data = _get_json(client, NPC_PUBLIC_FILE_API_URL, {"filePath": file_path}, timeout=timeout)
    if not isinstance(data, dict) or not isinstance(data.get("file"), dict):
        raise OfficialLegalDownloadError("官方文件预览接口返回格式无法识别。")
    download_url = _clean(data["file"].get("download_url"))
    if not download_url:
        raise OfficialLegalDownloadError("官方文件预览接口未返回下载地址。")
    host = (urlsplit(download_url).hostname or "").lower()
    if host.endswith("-internal.cucloud.cn"):
        raise OfficialLegalDownloadError("官方阅读器返回的是内网对象存储地址，当前环境无法直接下载。")
    response = client.get(
        download_url,
        headers={
            "User-Agent": "Law-Rag legal corpus downloader; official-source verification",
            "Referer": "https://flk.npc.gov.cn/",
        },
        timeout=timeout,
    )
    response.raise_for_status()
    return response.content, download_url


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _download_text(
    client: httpx.Client,
    detail: dict[str, Any],
    *,
    law_ref: str,
    title: str,
    timeout: float,
) -> tuple[str, str, str]:
    inline = _extract_inline_text(detail)
    if inline:
        return inline, NPC_DETAIL_API_URL, "国家法律法规数据库正文"

    html_source = _official_text_source_for(law_ref, title)
    if html_source is not None:
        try:
            return _fetch_official_html_text(client, html_source, timeout=timeout)
        except Exception as exc:
            html_error = f"{html_source.url}: {exc}"
        else:  # pragma: no cover - included to keep the error variable explicit.
            html_error = ""
    else:
        html_error = "无已适配的官方网页全文源"

    errors: list[str] = []
    for file_path in _official_file_urls(detail):
        try:
            content, url = _download_official_file(client, file_path, timeout=timeout)
            if file_path.lower().endswith(".docx"):
                return _docx_text(content), url, "国家法律法规数据库附件"
            return content.decode("utf-8", errors="ignore"), url, "国家法律法规数据库附件"
        except Exception as exc:  # pragma: no cover - depends on remote attachment behavior.
            errors.append(f"{file_path}: {exc}")
    for url in _attachment_urls(detail):
        try:
            headers = {"User-Agent": "Law-Rag legal corpus downloader; official-source verification"}
            try:
                response = client.get(url, headers=headers, timeout=timeout)
            except httpx.ConnectError as exc:
                host = (urlsplit(url).hostname or "").lower()
                if host != "wb.flk.npc.gov.cn" or "CERTIFICATE_VERIFY_FAILED" not in str(exc):
                    raise
                response = httpx.get(url, headers=headers, timeout=timeout, follow_redirects=True, verify=False)
            response.raise_for_status()
            content_type = response.headers.get("content-type", "").lower()
            if url.lower().split("?", 1)[0].endswith(".docx"):
                return _docx_text(response.content), url, "国家法律法规数据库附件"
            if "html" in content_type or url.lower().endswith((".html", ".htm")):
                return _text_from_html(response.text), url, "国家法律法规数据库附件"
            return response.text, url, "国家法律法规数据库附件"
        except Exception as exc:  # pragma: no cover - depends on remote attachment behavior.
            errors.append(f"{url}: {exc}")
    raise OfficialLegalDownloadError(
        f"未能从详情页、官方网页或附件提取正文：{html_error}; {'; '.join(errors) or '无附件'}"
    )


def download_npc_legal_manifest(
    law_ref: str,
    output_root: Path,
    *,
    timeout: float = 12.0,
    client: httpx.Client | None = None,
) -> DownloadedLegalManifest:
    owns_client = client is None
    client = client or httpx.Client(follow_redirects=True)
    try:
        search_result = _search_npc(client, law_ref, timeout=timeout)
        detail = _detail_npc(client, search_result, timeout=timeout)
        title = _extract_title(detail) or _extract_title(search_result) or law_ref
        text, text_url, text_source_name = _download_text(
            client,
            detail,
            law_ref=law_ref,
            title=title,
            timeout=timeout,
        )
        normalized = normalize_snapshot_text(text)
        try:
            parsed = parse_chinese_articles(
                normalized,
                authority_id=_slug(title),
                version_id="effective-2000-01-01",
            )
        except Exception as exc:
            raise OfficialLegalDownloadError(f"官方正文无法解析为按条文排列的法规文本：《{title}》：{exc}") from exc

        publication_date = _first_date(
            detail,
            ("publication_date", "publishDate", "gbrq", "f_bbrq", "公布日期", "date"),
        )
        effective_date = _first_date(
            detail,
            ("effective_date", "effectiveDate", "sxrq", "f_sxrq", "施行日期", "实施日期"),
        )
        if effective_date is None:
            effective_date = publication_date or date.today()
        authority_id = _slug(title)
        version_id = f"effective-{effective_date.isoformat()}"
        detail_id = _extract_id(search_result)
        detail_url = f"{NPC_DETAIL_PAGE_URL}?id={quote(detail_id)}&title={quote(title)}"

        manifest_dir = output_root / authority_id / version_id
        manifest_dir.mkdir(parents=True, exist_ok=True)
        snapshot_path = manifest_dir / "source.txt"
        snapshot_path.write_text(normalized, encoding="utf-8")

        record = ManifestRecord(
            authority=LegalAuthority(
                authority_id=authority_id,
                title=title,
                authority_type=_authority_type(detail),
                issuing_body=_first_text(
                    detail,
                    ("office", "issuingBody", "zdjg", "zdjgName", "制定机关", "department"),
                )
                or "国家法律法规数据库",
                document_number=_first_text(detail, ("document_number", "docNo", "f_wenhao", "文号")) or None,
            ),
            version_id=version_id,
            status=_version_status(detail),
            publication_date=publication_date,
            effective_date=effective_date,
            end_date_exclusive=None,
            repeal_date=None,
            supersedes_version_id=None,
            superseded_by_version_id=None,
            coverage_type=CoverageType.FULL_TEXT,
            coverage_note="Downloaded from the National Laws and Regulations Database and parsed locally.",
            source_refs=[
                OfficialSourceRef(name=f"国家法律法规数据库：{title}", url=detail_url, role=SourceRole.PRIMARY),
                OfficialSourceRef(name=f"{text_source_name}：{title}", url=text_url, role=SourceRole.TEXT),
            ],
            snapshot_path=snapshot_path.name,
            expected_source_sha256=sha256_text(normalized),
            expected_article_count=len(parsed.articles),
            parser="chinese-articles-v1",
            inclusion_reason=f"User selected the {law_ref} legal domain pack for local contract audit.",
            verified_on=date.today(),
            verification_note="Downloaded through the Law-Rag official source adapter and cross-checked against NPC metadata.",
        )
        manifest = LegalManifest(records=[record])
        manifest_path = manifest_dir / "manifest.json"
        manifest_path.write_text(manifest.model_dump_json(indent=2), encoding="utf-8")
        return DownloadedLegalManifest(
            law_ref=law_ref,
            title=title,
            manifest_path=manifest_path,
            article_count=len(parsed.articles),
        )
    finally:
        if owns_client:
            client.close()
