from __future__ import annotations

import hashlib
import io
import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas

from .report_export_models import ReportExportFormat, ReportExportResult
from .safe_persistence import atomic_write_text
from .tester_license import active_tester_watermark


class TesterReportWatermarkError(RuntimeError):
    pass


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _temp_path(destination: Path) -> Path:
    fd, name = tempfile.mkstemp(prefix=f".{destination.name}.tester-", suffix=".tmp", dir=destination.parent)
    os.close(fd)
    return Path(name)


def _watermark_docx(path: Path, watermark: str) -> None:
    document = Document(path)
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = watermark
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(7.5)

    temp = _temp_path(path)
    try:
        document.save(temp)
        if not temp.is_file() or temp.stat().st_size <= 0:
            raise TesterReportWatermarkError("DOCX tester watermark did not produce a non-empty document.")
        os.replace(temp, path)
    except Exception:
        temp.unlink(missing_ok=True)
        raise


def _overlay_pdf_page(width: float, height: float, watermark: str):
    buffer = io.BytesIO()
    overlay = canvas.Canvas(buffer, pagesize=(width, height))
    overlay.setFont("Helvetica", 7)
    overlay.setFillGray(0.48)
    overlay.drawCentredString(width / 2, 12, watermark)
    overlay.save()
    buffer.seek(0)
    reader = PdfReader(buffer)
    return reader.pages[0]


def _watermark_pdf(path: Path, watermark: str) -> None:
    reader = PdfReader(str(path))
    if not reader.pages:
        raise TesterReportWatermarkError("PDF tester watermark refused an empty report.")
    writer = PdfWriter()
    for page in reader.pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        page.merge_page(_overlay_pdf_page(width, height, watermark))
        writer.add_page(page)

    temp = _temp_path(path)
    try:
        with temp.open("wb") as handle:
            writer.write(handle)
            handle.flush()
            os.fsync(handle.fileno())
        if temp.stat().st_size <= 0:
            raise TesterReportWatermarkError("PDF tester watermark did not produce a non-empty document.")
        os.replace(temp, path)
    except Exception:
        temp.unlink(missing_ok=True)
        raise


def apply_tester_report_watermark(path: Path, result: ReportExportResult) -> ReportExportResult:
    """Add the active Tester ID to exported files and rebind export evidence.

    Normal/development builds have no tester-license requirement, so this helper
    becomes a no-op. The limited tester build applies the visible footer only
    after a valid signed license is active.
    """

    watermark = active_tester_watermark()
    if watermark is None:
        return result
    if result.format == ReportExportFormat.DOCX:
        _watermark_docx(path, watermark)
    elif result.format == ReportExportFormat.PDF:
        _watermark_pdf(path, watermark)
    else:
        raise TesterReportWatermarkError(f"Unsupported report format for tester watermark: {result.format}")

    updated = result.model_copy(
        update={
            "size_bytes": path.stat().st_size,
            "sha256": _sha256(path),
        }
    )
    manifest = path.parent / f"{path.name}.manifest.json"
    atomic_write_text(manifest, updated.model_dump_json(indent=2))
    return updated
