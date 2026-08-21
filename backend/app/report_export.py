from __future__ import annotations

import hashlib
import html
import json
import os
import tempfile
from pathlib import Path
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from .audit_planner import load_audit_plan
from .contract_structure import load_contract_structure
from .human_review import HumanReviewError, load_human_review
from .issue_legal_context import load_issue_legal_context
from .issue_primary_audit import load_issue_primary_audit
from .issue_review_report import load_issue_review_report
from .issue_secondary_review import load_issue_secondary_review
from .issue_workspace import load_issue_workspace_summary
from .report_export_models import (
    AuditReportDocument,
    ReportContractEvidence,
    ReportExportFormat,
    ReportExportResult,
    ReportHumanDecision,
    ReportIssue,
    ReportLegalEvidence,
)
from .storage import runtime_dir
from .workspace_models import WorkspaceOverallState


class ReportExportError(RuntimeError):
    pass


def _stable_sha256(payload: object) -> str:
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _source_span_index(contract) -> dict[str, ReportContractEvidence]:
    result: dict[str, ReportContractEvidence] = {}
    groups = (
        contract.title_candidates,
        contract.clauses,
        contract.unnumbered_blocks,
        contract.parties,
        contract.dates,
        contract.money_mentions,
        contract.percentages,
        contract.identifiers,
        contract.references,
        contract.structured_blocks,
    )
    for group in groups:
        for item in group:
            for span in getattr(item, "source_spans", []):
                source_method = getattr(span.source_method, "value", str(span.source_method))
                for evidence_id in span.evidence_ids:
                    result.setdefault(
                        evidence_id,
                        ReportContractEvidence(
                            evidence_id=evidence_id,
                            quote=span.quote or None,
                            page_number=span.page_number,
                            source_method=source_method,
                        ),
                    )
    return result


def _assert_exact_issue_coverage(plan, legal, primary, secondary, review) -> None:
    expected = {item.issue_id for item in plan.issues}
    actual_sets = {
        "issue-legal-context": {item.issue_id for item in legal.issues},
        "issue-primary-audit": {item.issue_id for item in primary.results},
        "issue-secondary-review": {item.issue_id for item in secondary.results},
        "issue-review-report": {item.issue_id for item in review.comparisons},
    }
    for label, actual in actual_sets.items():
        if actual != expected:
            missing = sorted(expected - actual)
            extra = sorted(actual - expected)
            raise ReportExportError(
                f"{label} does not exactly cover the AuditPlan issues; missing={missing}, extra={extra}."
            )


def build_audit_report(job_id: UUID) -> AuditReportDocument:
    workspace = load_issue_workspace_summary(job_id)
    if workspace.overall_state not in {
        WorkspaceOverallState.COMPLETE,
        WorkspaceOverallState.HUMAN_REVIEW_REQUIRED,
    }:
        raise ReportExportError(
            f"A report may only be exported from a validated ISSUE_V1 comparison state; current state={workspace.overall_state.value}."
        )
    if workspace.document is None or workspace.coverage is None:
        raise ReportExportError("Validated document and AuditPlan coverage metadata are required for report export.")

    contract = load_contract_structure(job_id)
    plan = load_audit_plan(job_id)
    legal = load_issue_legal_context(job_id)
    primary = load_issue_primary_audit(job_id)
    secondary = load_issue_secondary_review(job_id)
    review = load_issue_review_report(job_id)
    _assert_exact_issue_coverage(plan, legal, primary, secondary, review)

    try:
        human = load_human_review(job_id)
    except (FileNotFoundError, HumanReviewError):
        human = None

    contract_evidence = _source_span_index(contract)
    legal_by_id = {item.issue_id: item for item in legal.issues}
    primary_by_id = {item.issue_id: item for item in primary.results}
    secondary_by_id = {item.issue_id: item for item in secondary.results}
    comparison_by_id = {item.issue_id: item for item in review.comparisons}
    human_by_id = {}
    if human is not None:
        human_by_id = {
            key.removeprefix("issue:"): value
            for key, value in human.latest_by_target.items()
            if key.startswith("issue:")
        }

    issues: list[ReportIssue] = []
    for plan_issue in plan.issues:
        legal_item = legal_by_id[plan_issue.issue_id]
        primary_item = primary_by_id[plan_issue.issue_id]
        secondary_item = secondary_by_id[plan_issue.issue_id]
        comparison = comparison_by_id[plan_issue.issue_id]
        human_item = human_by_id.get(plan_issue.issue_id)

        evidence_ids = sorted(
            set(plan_issue.contract_evidence_ids)
            | set(primary_item.contract_evidence_ids)
            | set(secondary_item.contract_evidence_ids)
        )
        contract_items = [
            contract_evidence.get(evidence_id, ReportContractEvidence(evidence_id=evidence_id))
            for evidence_id in evidence_ids
        ]

        legal_items = [
            ReportLegalEvidence(
                legal_evidence_id=hit.legal_evidence_id,
                authority_id=hit.candidate.authority_id,
                authority_title=hit.candidate.authority_title,
                version_id=hit.candidate.version_id,
                article_token=hit.candidate.article_token,
                article_text=hit.candidate.article_text,
                effective_date=hit.candidate.effective_date.isoformat(),
                end_date_exclusive=(
                    hit.candidate.end_date_exclusive.isoformat()
                    if hit.candidate.end_date_exclusive is not None
                    else None
                ),
                coverage_type=hit.candidate.coverage_type,
            )
            for hit in legal_item.legal_evidence
        ]

        human_decision = None
        if human_item is not None:
            human_decision = ReportHumanDecision(
                state=human_item.state.value,
                revision=human_item.revision,
                decided_at=human_item.decided_at,
                reviewer_note=human_item.reviewer_note,
                is_stale=human_item.is_stale,
            )

        issues.append(
            ReportIssue(
                issue_id=plan_issue.issue_id,
                topic=plan_issue.topic,
                priority=plan_issue.priority.value,
                why_review=list(plan_issue.why_review),
                questions=list(plan_issue.questions),
                primary_state=primary_item.state.value,
                primary_severity=primary_item.severity.value,
                primary_title=primary_item.title,
                primary_reasoning=primary_item.reasoning_summary,
                primary_suggestion=primary_item.suggestion,
                primary_evidence_sufficiency=primary_item.evidence_sufficiency.value,
                secondary_assessment=secondary_item.assessment.value,
                secondary_coverage_assessment=secondary_item.coverage_assessment.value,
                secondary_severity=secondary_item.severity.value,
                secondary_reasoning=secondary_item.reasoning_summary,
                secondary_suggestion=secondary_item.suggestion,
                comparison_state=comparison.overall_state.value,
                requires_human_review=comparison.requires_human_review,
                comparison_reasons=list(comparison.reasons),
                omission_title=comparison.omission_title,
                omission_reasoning=comparison.omission_reasoning,
                contract_evidence=contract_items,
                legal_evidence=legal_items,
                human_decision=human_decision,
            )
        )

    source_fingerprints = {
        "contract_source": contract.source_fingerprint,
        "audit_plan_input": plan.planner_input_fingerprint,
        "issue_legal_context": legal.artifact_fingerprint,
        "issue_primary_audit": primary.artifact_fingerprint,
        "issue_secondary_review": secondary.artifact_fingerprint,
        "issue_review_report": review.artifact_fingerprint,
    }

    payload = {
        "job_id": str(job_id),
        "filename": workspace.document.filename,
        "document_kind": workspace.document.document_kind,
        "as_of": review.as_of,
        "overall_state": workspace.overall_state.value,
        "contract_type": workspace.coverage.contract_type.value,
        "planning_mode": workspace.coverage.planning_mode.value,
        "planning_coverage_complete": workspace.coverage.coverage_complete,
        "canonical_object_count": workspace.coverage.canonical_object_count,
        "reviewed_with_issue_count": workspace.coverage.reviewed_with_issue_count,
        "reviewed_no_specific_issue_count": workspace.coverage.reviewed_no_specific_issue_count,
        "primary_provider": review.primary_provider,
        "primary_model": review.primary_model,
        "secondary_provider": review.secondary_provider,
        "secondary_model": review.secondary_model,
        "final_review_state": review.final_state.value,
        "human_review_required_count": workspace.review.human_review_required_count,
        "human_review_resolved_required_count": workspace.review.human_review_resolved_required_count,
        "human_review_outstanding_required_count": workspace.review.human_review_outstanding_required_count,
        "issues": [item.model_dump(mode="json") for item in issues],
        "source_uncertainty": sorted(set(workspace.source_uncertainty)),
        "warnings": sorted(set(workspace.warnings + review.warnings)),
        "source_fingerprints": source_fingerprints,
    }
    return AuditReportDocument(
        **payload,
        report_content_fingerprint=_stable_sha256(payload),
    )


def _safe_text(value: str | None) -> str:
    return (value or "—").strip() or "—"


def _docx_add_bullets(document: Document, values: list[str]) -> None:
    if not values:
        document.add_paragraph("—")
        return
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _render_docx(report: AuditReportDocument, destination: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "Microsoft YaHei"

    title = document.add_heading("Law-Rag 合同审计报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status = document.add_paragraph(
        "状态：" + (
            "需要人工复核，本文档不是最终法律意见"
            if report.human_review_outstanding_required_count
            else "审计链完整；仍应由有资格人员结合实际事实作最终判断"
        )
    )
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in (
        ("源文件", report.filename),
        ("Job ID", str(report.job_id)),
        ("审计基准日", report.as_of),
        ("合同类型", report.contract_type),
        ("规划模式", report.planning_mode),
        ("规划覆盖", f"{report.canonical_object_count} 个 canonical object；complete={report.planning_coverage_complete}"),
        ("Primary", f"{report.primary_provider} / {report.primary_model}"),
        ("Secondary", f"{report.secondary_provider} / {report.secondary_model}"),
        ("人工复核", f"{report.human_review_resolved_required_count}/{report.human_review_required_count} 已完成"),
        ("内容指纹", report.report_content_fingerprint),
    ):
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)

    document.add_heading("审计结论与证据", level=1)
    for index, issue in enumerate(report.issues, start=1):
        document.add_heading(f"{index}. {issue.topic} [{issue.priority}]", level=2)
        document.add_paragraph(f"Issue ID：{issue.issue_id}")
        document.add_paragraph(f"对比状态：{issue.comparison_state}；需要人工复核：{'是' if issue.requires_human_review else '否'}")
        document.add_paragraph("为什么审：")
        _docx_add_bullets(document, issue.why_review)
        document.add_paragraph("Primary 结论：")
        document.add_paragraph(
            f"{issue.primary_title} | {issue.primary_state} | severity={issue.primary_severity} | evidence={issue.primary_evidence_sufficiency}"
        )
        document.add_paragraph(issue.primary_reasoning)
        document.add_paragraph(f"建议：{issue.primary_suggestion}")
        document.add_paragraph("Secondary 复核：")
        document.add_paragraph(
            f"assessment={issue.secondary_assessment} | coverage={issue.secondary_coverage_assessment} | severity={issue.secondary_severity}"
        )
        document.add_paragraph(issue.secondary_reasoning)
        document.add_paragraph(f"建议：{issue.secondary_suggestion}")
        if issue.comparison_reasons:
            document.add_paragraph("确定性对比理由：")
            _docx_add_bullets(document, issue.comparison_reasons)
        if issue.omission_title or issue.omission_reasoning:
            document.add_paragraph(f"可能遗漏：{_safe_text(issue.omission_title)}")
            document.add_paragraph(_safe_text(issue.omission_reasoning))

        document.add_paragraph("合同证据：")
        if issue.contract_evidence:
            for evidence in issue.contract_evidence:
                where = f"第 {evidence.page_number} 页" if evidence.page_number else "结构锚点"
                document.add_paragraph(
                    f"{evidence.evidence_id} · {where} · {_safe_text(evidence.source_method)}\n{_safe_text(evidence.quote)}",
                    style="List Bullet",
                )
        else:
            document.add_paragraph("—")

        document.add_paragraph("法律依据：")
        if issue.legal_evidence:
            for evidence in issue.legal_evidence:
                validity = evidence.effective_date + (f" 至 {evidence.end_date_exclusive}（不含）" if evidence.end_date_exclusive else " 起")
                document.add_paragraph(
                    f"{evidence.authority_title} {evidence.article_token} · version={evidence.version_id} · {validity}\n{evidence.article_text}",
                    style="List Bullet",
                )
        else:
            document.add_paragraph("本地法律语料未提供可引用证据。")

        document.add_paragraph("人工决定：")
        if issue.human_decision is None:
            document.add_paragraph("尚无人工决定。")
        else:
            decision = issue.human_decision
            document.add_paragraph(
                f"{decision.state} · revision={decision.revision} · {decision.decided_at.isoformat()} · stale={'YES' if decision.is_stale else 'NO'}"
            )
            if decision.reviewer_note:
                document.add_paragraph(decision.reviewer_note)

    document.add_heading("不确定性与警告", level=1)
    _docx_add_bullets(document, [*report.source_uncertainty, *report.warnings])
    document.add_heading("来源指纹", level=1)
    for label, fingerprint in sorted(report.source_fingerprints.items()):
        document.add_paragraph(f"{label}: {fingerprint}")

    props = document.core_properties
    props.title = "Law-Rag 合同审计报告"
    props.subject = f"Job {report.job_id}"
    props.author = "Law-Rag"
    document.save(destination)


def _pdf_paragraph(text: str, style) -> Paragraph:
    escaped = html.escape(text).replace("\n", "<br/>")
    return Paragraph(escaped, style)


def _render_pdf(report: AuditReportDocument, destination: Path) -> None:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except KeyError:
        pass

    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "LawRagBody",
        parent=base["BodyText"],
        fontName="STSong-Light",
        fontSize=9.5,
        leading=14,
        spaceAfter=5,
    )
    heading = ParagraphStyle(
        "LawRagHeading",
        parent=body,
        fontSize=15,
        leading=20,
        spaceBefore=10,
        spaceAfter=7,
    )
    title_style = ParagraphStyle(
        "LawRagTitle",
        parent=heading,
        fontSize=22,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    small = ParagraphStyle("LawRagSmall", parent=body, fontSize=8, leading=11)

    doc = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Law-Rag 合同审计报告",
        author="Law-Rag",
    )
    story = [_pdf_paragraph("Law-Rag 合同审计报告", title_style)]
    status_text = (
        "状态：需要人工复核，本文档不是最终法律意见"
        if report.human_review_outstanding_required_count
        else "状态：审计链完整；仍应由有资格人员结合实际事实作最终判断"
    )
    story += [
        _pdf_paragraph(status_text, body),
        _pdf_paragraph(f"源文件：{report.filename}", body),
        _pdf_paragraph(f"Job ID：{report.job_id}", small),
        _pdf_paragraph(f"审计基准日：{report.as_of}　合同类型：{report.contract_type}　规划模式：{report.planning_mode}", body),
        _pdf_paragraph(
            f"规划覆盖：{report.canonical_object_count} 个 canonical object；complete={report.planning_coverage_complete}　人工复核：{report.human_review_resolved_required_count}/{report.human_review_required_count}",
            body,
        ),
        _pdf_paragraph(f"内容指纹：{report.report_content_fingerprint}", small),
        Spacer(1, 4 * mm),
    ]

    for index, issue in enumerate(report.issues, start=1):
        story.append(_pdf_paragraph(f"{index}. {issue.topic} [{issue.priority}]", heading))
        story.append(_pdf_paragraph(f"Issue ID：{issue.issue_id}", small))
        story.append(_pdf_paragraph(
            f"对比状态：{issue.comparison_state}；需要人工复核：{'是' if issue.requires_human_review else '否'}",
            body,
        ))
        for value in issue.why_review:
            story.append(_pdf_paragraph(f"• {value}", body))
        story.append(_pdf_paragraph(
            f"Primary：{issue.primary_title} | {issue.primary_state} | severity={issue.primary_severity} | evidence={issue.primary_evidence_sufficiency}",
            body,
        ))
        story.append(_pdf_paragraph(issue.primary_reasoning, body))
        story.append(_pdf_paragraph(f"Primary 建议：{issue.primary_suggestion}", body))
        story.append(_pdf_paragraph(
            f"Secondary：assessment={issue.secondary_assessment} | coverage={issue.secondary_coverage_assessment} | severity={issue.secondary_severity}",
            body,
        ))
        story.append(_pdf_paragraph(issue.secondary_reasoning, body))
        story.append(_pdf_paragraph(f"Secondary 建议：{issue.secondary_suggestion}", body))
        for reason in issue.comparison_reasons:
            story.append(_pdf_paragraph(f"确定性对比：• {reason}", body))
        if issue.omission_title or issue.omission_reasoning:
            story.append(_pdf_paragraph(f"可能遗漏：{_safe_text(issue.omission_title)}", body))
            story.append(_pdf_paragraph(_safe_text(issue.omission_reasoning), body))

        story.append(_pdf_paragraph("合同证据", heading))
        if issue.contract_evidence:
            for evidence in issue.contract_evidence:
                where = f"第 {evidence.page_number} 页" if evidence.page_number else "结构锚点"
                story.append(_pdf_paragraph(
                    f"• {evidence.evidence_id} · {where} · {_safe_text(evidence.source_method)}<br/>{_safe_text(evidence.quote)}",
                    body,
                ))
        else:
            story.append(_pdf_paragraph("—", body))

        story.append(_pdf_paragraph("法律依据", heading))
        if issue.legal_evidence:
            for evidence in issue.legal_evidence:
                validity = evidence.effective_date + (f" 至 {evidence.end_date_exclusive}（不含）" if evidence.end_date_exclusive else " 起")
                story.append(_pdf_paragraph(
                    f"• {evidence.authority_title} {evidence.article_token} · version={evidence.version_id} · {validity}<br/>{evidence.article_text}",
                    body,
                ))
        else:
            story.append(_pdf_paragraph("本地法律语料未提供可引用证据。", body))

        if issue.human_decision is None:
            story.append(_pdf_paragraph("人工决定：尚无人工决定。", body))
        else:
            decision = issue.human_decision
            story.append(_pdf_paragraph(
                f"人工决定：{decision.state} · revision={decision.revision} · {decision.decided_at.isoformat()} · stale={'YES' if decision.is_stale else 'NO'}",
                body,
            ))
            if decision.reviewer_note:
                story.append(_pdf_paragraph(f"人工备注：{decision.reviewer_note}", body))
        if index < len(report.issues):
            story.append(PageBreak())

    story.append(PageBreak())
    story.append(_pdf_paragraph("不确定性与警告", heading))
    values = [*report.source_uncertainty, *report.warnings]
    if values:
        for value in values:
            story.append(_pdf_paragraph(f"• {value}", body))
    else:
        story.append(_pdf_paragraph("—", body))
    story.append(_pdf_paragraph("来源指纹", heading))
    for label, fingerprint in sorted(report.source_fingerprints.items()):
        story.append(_pdf_paragraph(f"{label}: {fingerprint}", small))

    doc.build(story)


def _atomic_render(destination: Path, renderer) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, name = tempfile.mkstemp(prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent)
    os.close(fd)
    temp = Path(name)
    try:
        renderer(temp)
        if not temp.is_file() or temp.stat().st_size <= 0:
            raise ReportExportError("Report renderer did not produce a non-empty file.")
        with temp.open("rb") as handle:
            os.fsync(handle.fileno())
        os.replace(temp, destination)
    except Exception:
        temp.unlink(missing_ok=True)
        raise


def export_audit_report(job_id: UUID, format: ReportExportFormat) -> tuple[Path, ReportExportResult]:
    report = build_audit_report(job_id)
    export_dir = runtime_dir() / "exports" / str(job_id)
    filename = f"Law-Rag-Audit-{str(job_id)[:8]}.{format.value}"
    destination = export_dir / filename
    if destination.is_symlink() or export_dir.is_symlink():
        raise ReportExportError("Report export path must not be a symlink.")

    renderer = _render_docx if format == ReportExportFormat.DOCX else _render_pdf
    _atomic_render(destination, lambda path: renderer(report, path))
    result = ReportExportResult(
        job_id=job_id,
        format=format,
        filename=filename,
        size_bytes=destination.stat().st_size,
        sha256=_file_sha256(destination),
        report_content_fingerprint=report.report_content_fingerprint,
    )
    manifest = export_dir / f"{filename}.manifest.json"
    manifest.write_text(result.model_dump_json(indent=2), encoding="utf-8")
    return destination, result
